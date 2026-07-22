import os
import subprocess
import time
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Error: Falta la librería python-docx. Ejecuta: python -m pip install python-docx")
    exit()

def agregar_codigo_consola(doc, texto):
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(12)

def insertar_imagen(doc, ruta_imagen, titulo):
    if os.path.exists(ruta_imagen):
        p = doc.add_paragraph()
        run = p.add_run()
        run.add_picture(ruta_imagen, width=Inches(5.8))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        p_cap = doc.add_paragraph(titulo)
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.runs[0].italic = True
        p_cap.runs[0].font.size = Pt(9.5)
        p_cap.paragraph_format.space_after = Pt(12)
    else:
        doc.add_paragraph(f"[Nota: Imagen no encontrada en la ruta: {ruta_imagen}]")

def generar_informe_completo():
    print("\n=== GENERANDO INFORME WORD COMPLETO CON EVIDENCIAS REALES ===")
    
    doc = Document()
    
    # --- Portada Formal ---
    titulo = doc.add_heading('Informe Técnico de Práctica:\nOrquestación de Smart Contract en Kubernetes', 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo.paragraph_format.space_after = Pt(24)

    p_aut = doc.add_paragraph()
    p_aut.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_aut = p_aut.add_run('Autor: Andrés Alejandro Silva Aguilar\n')
    r_aut.bold = True
    r_aut.font.size = Pt(12)
    
    p_aut.add_run('Institución: Universidad Técnica Particular de Loja (UTPL)\n')
    p_aut.add_run('Programa: Especialización en Arquitectura Cloud y Blockchain\n')
    p_aut.add_run('Fecha de Ejecución: Julio 2026\n')
    
    p_git = doc.add_paragraph()
    p_git.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_git.add_run('Repositorio GitHub Oficial:\n').bold = True
    p_git.add_run('https://github.com/aalejos3626/smartcontract_kubernetes_andres_silva')
    
    doc.add_page_break()

    # --- Rutas de las imágenes corregidas con r'...' ---
    path_img_repo = r"C:\Users\AlejandroSilva\Documents\Blockchain\MIGRACION Y ADMIN DE SERVICIO\Kubernetes\image_1a3fab.png"
    path_img_services = r"C:\Users\AlejandroSilva\Documents\Blockchain\MIGRACION Y ADMIN DE SERVICIO\Kubernetes\image_1a38c1.png"
    path_img_replicas = r"C:\Users\AlejandroSilva\Documents\Blockchain\MIGRACION Y ADMIN DE SERVICIO\Kubernetes\image_1a384a.png"

    # --- Sección 1 ---
    doc.add_heading('1. Introducción y Objetivos del Smart Contract', level=1)
    doc.add_paragraph('El presente informe documenta la implementación y orquestación de un contrato inteligente desarrollado en Solidity (SmartContract_AndresSilva.sol), desplegado sobre un clúster de Kubernetes local utilizando Docker Desktop.')
    doc.add_paragraph('Objetivos principales:', style='List Bullet')
    doc.add_paragraph('Garantizar alta disponibilidad mediante un controlador ReplicaSet y despliegue automatizado.', style='List Bullet 2')
    doc.add_paragraph('Validar operaciones de escalado horizontal en caliente y actualizaciones continuas sin tiempo de inactividad (Rolling Updates).', style='List Bullet 2')

    # --- Sección 2 ---
    doc.add_heading('2. Arquitectura de Despliegue y Diagrama', level=1)
    doc.add_paragraph('La solución se apoya en un diseño desacoplado donde un objeto de tipo Service (ClusterIP) expone el puerto estándar de EVM (8545) hacia el clúster, conectándose directamente al Deployment que administra las réplicas del nodo Ganache CLI que aloja el contrato inteligente.')
    doc.add_paragraph('Nota de Diseño: La arquitectura completa se encuentra modelada de forma vectorial y editable en el archivo "arquitectura_kubernetes_andres_silva.drawio" almacenado en la raíz del repositorio oficial.')
    
    insertar_imagen(doc, path_img_repo, "Figura 1: Evidencia del repositorio GitHub sincronizado con los artefactos de la práctica.")

    # --- Sección 3 ---
    doc.add_heading('3. Proceso de Orquestación y Evidencias Prácticas', level=1)
    
    doc.add_heading('A. Despliegue Inicial y Verificación de Nodos/Servicios', level=2)
    doc.add_paragraph('Se aplicó el manifiesto YAML inicial configurando 3 réplicas iniciales para el contrato inteligente.')
    agregar_codigo_consola(doc, "> kubectl apply -f k8s/deployment_blockchain.yaml\n> kubectl get pods\n> kubectl get services")
    
    insertar_imagen(doc, path_img_services, "Figura 2: Panel de Docker Desktop visualizando el Control Plane, los pods activos y el servicio ClusterIP expuesto en el puerto 8545.")

    doc.add_heading('B. Escalado Horizontal a 5 Réplicas', level=2)
    doc.add_paragraph('Mediante el comando de escalado horizontal, se incrementó la capacidad de procesamiento de nodos a 5 instancias simultáneas en caliente:')
    agregar_codigo_consola(doc, "> kubectl scale deployment/blockchain-contract-deployment --replicas=5\n> kubectl get pods")
    
    insertar_imagen(doc, path_img_replicas, "Figura 3: Comprobación gráfica en Docker Desktop de las 5 réplicas ejecutándose correctamente en estado Running.")

    doc.add_heading('C. Actualización Continua (Rolling Update)', level=2)
    doc.add_paragraph('Se ejecutó la actualización controlada de la imagen del contenedor para simular un pase a producción con una nueva versión del entorno EVM sin interrumpir el servicio operativo:')
    agregar_codigo_consola(doc, "> kubectl set image deployment/blockchain-contract-deployment node-container=trufflesuite/ganache-cli:v6.12.2\n> kubectl rollout status deployment/blockchain-contract-deployment")

    nombre_archivo = "Informe_Orquestacion_Blockchain_Andres_Silva.docx"
    doc.save(nombre_archivo)
    print(f"\n=== ¡DOCUMENTO GENERADO CON ÉXITO! ===")
    print(f"Archivo guardado como: {nombre_archivo}")

if __name__ == "__main__":
    generar_informe_completo()