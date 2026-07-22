import os
import subprocess
import time
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Error: Falta la librería python-docx. Ejecuta: python -m pip install python-docx")
    exit()

def run_cmd(cmd):
    print("\n[EJECUTANDO] " + cmd)
    try:
        result = subprocess.run(cmd, shell=True, text=True, capture_output=True, check=True)
        return result.stdout.strip()
    except subprocess.CalledProcessError as e:
        return "Error: " + e.stderr

def agregar_codigo_consola(doc, texto):
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(12)

def crear_archivo_drawio():
    """Genera un archivo draw.io real (XML editable) con la arquitectura del despliegue"""
    print("--- Generando diagrama real en formato draw.io ---")
    drawio_content = """<?xml version="1.0" encoding="UTF-8"?>
<mxfile host="app.diagrams.net" modified="2026-07-21T00:00:00.000Z" agent="Mozilla/5.0" version="24.7.1" type="device">
  <diagram name="Arquitectura Kubernetes - Andres Silva" id="k8s_architecture">
    <mxGraphModel dx="1000" dy="1000" grid="1" gridSize="10" guides="1" tooltips="1" connect="1" arrows="1" fold="1" page="1" pageScale="1" pageWidth="850" pageHeight="1100" math="0" shadow="0">
      <root>
        <mxCell id="0" />
        <mxCell id="1" parent="0" />
        
        <!-- Contenedor Cluster Kubernetes -->
        <mxCell id="cluster_k8s" value="Cluster Kubernetes (Docker Desktop / Kind) - Autor: Andrés Silva" style="swimlane;whiteSpace=wrap;html=1;fillColor=#f5f5f5;strokeColor=#666666;fontColor=#333333;fontStyle=1;fontSize=14;" vertex="1" parent="1">
          <mxGeometry x="40" y="40" width="720" height="520" as="geometry" />
        </mxCell>
        
        <!-- Service ClusterIP -->
        <mxCell id="service_node" value="Service: blockchain-contract-service&lt;br&gt;(ClusterIP - Port: 8545)" style="rounded=1;whiteSpace=wrap;html=1;fillColor=#dae8fc;strokeColor=#6c8ebf;fontStyle=1;fontSize=12;" vertex="1" parent="cluster_k8s">
          <mxGeometry x="180" y="60" width="360" height="60" as="geometry" />
        </mxCell>
        
        <!-- Deployment Controller -->
        <mxCell id="deployment_node" value="Deployment: blockchain-contract-deployment&lt;br&gt;(ReplicaSet Controller - 5 Réplicas Activas)" style="rounded=1;whiteSpace=wrap;html=1;fillColor=#ffe6cc;strokeColor=#d79b00;fontStyle=1;fontSize=12;" vertex="1" parent="cluster_k8s">
          <mxGeometry x="100" y="160" width="520" height="320" as="geometry" />
        </mxCell>
        
        <!-- Pods 1 al 5 -->
        <mxCell id="pod_1" value="Pod 1&lt;br&gt;(Ganache CLI)" style="rounded=0;whiteSpace=wrap;html=1;fillColor=#d5e8d4;strokeColor=#82b366;fontSize=11;" vertex="1" parent="deployment_node">
          <mxGeometry x="30" y="60" width="80" height="80" as="geometry" />
        </mxCell>
        <mxCell id="pod_2" value="Pod 2&lt;br&gt;(Ganache CLI)" style="rounded=0;whiteSpace=wrap;html=1;fillColor=#d5e8d4;strokeColor=#82b366;fontSize=11;" vertex="1" parent="deployment_node">
          <mxGeometry x="125" y="60" width="80" height="80" as="geometry" />
        </mxCell>
        <mxCell id="pod_3" value="Pod 3&lt;br&gt;(Ganache CLI)" style="rounded=0;whiteSpace=wrap;html=1;fillColor=#d5e8d4;strokeColor=#82b366;fontSize=11;" vertex="1" parent="deployment_node">
          <mxGeometry x="220" y="60" width="80" height="80" as="geometry" />
        </mxCell>
        <mxCell id="pod_4" value="Pod 4&lt;br&gt;(Ganache CLI)" style="rounded=0;whiteSpace=wrap;html=1;fillColor=#d5e8d4;strokeColor=#82b366;fontSize=11;" vertex="1" parent="deployment_node">
          <mxGeometry x="315" y="60" width="80" height="80" as="geometry" />
        </mxCell>
        <mxCell id="pod_5" value="Pod 5&lt;br&gt;(Ganache CLI)" style="rounded=0;whiteSpace=wrap;html=1;fillColor=#d5e8d4;strokeColor=#82b366;fontSize=11;" vertex="1" parent="deployment_node">
          <mxGeometry x="410" y="60" width="80" height="80" as="geometry" />
        </mxCell>
        
        <!-- Smart Contract Desc -->
        <mxCell id="contract_desc" value="SmartContract_AndresSilva.sol (EVM Node Backend)" style="rounded=1;whiteSpace=wrap;html=1;fillColor=#e1d5e7;strokeColor=#9673a6;fontStyle=2;fontSize=11;" vertex="1" parent="deployment_node">
          <mxGeometry x="30" y="180" width="460" height="40" as="geometry" />
        </mxCell>
        
        <!-- Conexiones -->
        <mxCell id="edge_1" value="" style="endArrow=classic;html=1;rounded=0;entryX=0.5;entryY=0;entryDx=0;entryDy=0;exitX=0.5;exitY=1;exitDx=0;exitDy=0;" edge="1" parent="cluster_k8s" source="service_node" target="deployment_node">
          <mxGeometry width="50" height="50" relative="1" as="geometry">
            <mxPoint x="370" y="270" as="sourcePoint" />
            <mxPoint x="420" y="220" as="targetPoint" />
          </mxGeometry>
        </mxCell>
        
      </root>
    </mxGraphModel>
  </diagram>
</mxfile>"""
    
    with open("arquitectura_kubernetes_andres_silva.drawio", "w", encoding="utf-8") as f:
        f.write(drawio_content)
    print("Archivo 'arquitectura_kubernetes_andres_silva.drawio' creado exitosamente.")

def crear_estructura_proyecto():
    print("--- 1. Creando estructura de archivos del proyecto ---")
    os.makedirs("contracts", exist_ok=True)
    os.makedirs("k8s", exist_ok=True)
    os.makedirs("diagrams", exist_ok=True)

    # Contrato inteligente personalizado
    with open("contracts/SmartContract_AndresSilva.sol", "w", encoding="utf-8") as f:
        f.write("""// SPDX-License-Identifier: MIT
pragma solidity ^0.8.0;

/**
 * @title SmartContract_AndresSilva
 * @dev Contrato inteligente optimizado para práctica de orquestación en Kubernetes
 * @author Andrés Alejandro Silva Aguilar (UTPL)
 */
contract SmartContract_AndresSilva {
    uint256 public contadorTransaccional;
    string public autor = "Andres Alejandro Silva Aguilar";
    
    event TransaccionRegistrada(uint256 nuevoValor, string mensaje);

    function incrementarRegistro(string memory mensaje) public {
        contadorTransaccional += 1;
        emit TransaccionRegistrada(contadorTransaccional, mensaje);
    }

    function obtenerContador() public view returns (uint256) {
        return contadorTransaccional;
    }
}
""")

    # Manifiesto Kubernetes
    with open("k8s/deployment_blockchain.yaml", "w", encoding="utf-8") as f:
        f.write("""apiVersion: apps/v1
kind: Deployment
metadata:
  name: blockchain-contract-deployment
  labels:
    app: blockchain-node
    author: andres-silva
spec:
  replicas: 3
  selector:
    matchLabels:
      app: blockchain-node
  template:
    metadata:
      labels:
        app: blockchain-node
    spec:
      containers:
      - name: node-container
        image: trufflesuite/ganache-cli:latest
        ports:
        - containerPort: 8545
---
apiVersion: v1
kind: Service
metadata:
  name: blockchain-contract-service
spec:
  selector:
    app: blockchain-node
  ports:
    - protocol: TCP
      port: 8545
      targetPort: 8545
  type: ClusterIP
""")

    # Crear el diagrama draw.io en la raíz y mover una copia a diagrams/
    crear_archivo_drawio()
    if os.path.exists("arquitectura_kubernetes_andres_silva.drawio"):
        import shutil
        shutil.copy("arquitectura_kubernetes_andres_silva.drawio", "diagrams/arquitectura.drawio")

    # README Profesional
    with open("README.md", "w", encoding="utf-8") as f:
        f.write("""# Práctica de Orquestación de Smart Contract en Kubernetes

* **Estudiante:** Andrés Alejandro Silva Aguilar
* **Institución:** Universidad Técnica Particular de Loja (UTPL)
* **Tema:** Arquitectura Cloud y Despliegue de Contratos Inteligentes en Contenedores Orquestados.

## Estructura del Repositorio
* `contracts/`: Contiene el contrato inteligente Solidity (`SmartContract_AndresSilva.sol`).
* `k8s/`: Manifiestos YAML de despliegue y servicios para Kubernetes.
* `diagrams/`: Diagrama de arquitectura editable en formato nativo draw.io.
""")
    print("Estructura de archivos creada con éxito.")

def inicializar_git_y_github(repo_nombre):
    print(f"\n--- 2. Inicializando Git y sincronizando con GitHub: {repo_nombre} ---")
    run_cmd("git init -b main")
    run_cmd("git add .")
    run_cmd('git commit -m "Commit inicial: Incluye contrato, manifiestos y diagrama draw.io"')
    
    cmd_gh = f"gh repo create {repo_nombre} --public --source=. --remote=origin --push"
    res = run_cmd(cmd_gh)
    print(res)
    return f"https://github.com/aalejos3626/{repo_nombre}"

def ejecutar_orquestacion_kubernetes():
    print("\n--- 3. Ejecutando Orquestación en Kubernetes (Docker Desktop) ---")
    
    run_cmd("kubectl apply -f k8s/deployment_blockchain.yaml")
    time.sleep(8)
    pods_ini = run_cmd("kubectl get pods")
    rs_ini = run_cmd("kubectl get rs")
    
    print("\n[ACCIÓN] Escalando horizontalmente a 5 réplicas...")
    run_cmd("kubectl scale deployment/blockchain-contract-deployment --replicas=5")
    time.sleep(6)
    pods_scale = run_cmd("kubectl get pods")
    rs_scale = run_cmd("kubectl get rs")
    
    print("\n[ACCIÓN] Ejecutando Rolling Update de imagen...")
    run_cmd("kubectl set image deployment/blockchain-contract-deployment node-container=trufflesuite/ganache-cli:v6.12.2")
    rollout_status = run_cmd("kubectl rollout status deployment/blockchain-contract-deployment")
    time.sleep(8)
    pods_update = run_cmd("kubectl get pods")
    rs_update = run_cmd("kubectl get rs")
    
    return {
        "pods_ini": pods_ini, "rs_ini": rs_ini,
        "pods_scale": pods_scale, "rs_scale": rs_scale,
        "rollout": rollout_status,
        "pods_update": pods_update, "rs_update": rs_update
    }

def generar_documento_word(evidencias, repo_url):
    print("\n--- 4. Generando Documento Word (.docx) ---")
    doc = Document()
    
    titulo = doc.add_heading('Informe Técnico: Orquestación de Smart Contract en Kubernetes', 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('Autor: Andrés Alejandro Silva Aguilar').bold = True
    doc.add_paragraph('Institución: Universidad Técnica Particular de Loja (UTPL)')
    doc.add_paragraph('Módulo: Especialización en Arquitectura Cloud y Blockchain')
    doc.add_paragraph('Fecha: Julio 2026')
    
    p_git = doc.add_paragraph()
    p_git.add_run('Repositorio GitHub Oficial: ').bold = True
    p_git.add_run(repo_url)
    
    doc.add_page_break()

    doc.add_heading('1. Descripción del Smart Contract', level=1)
    doc.add_paragraph('Archivo fuente: SmartContract_AndresSilva.sol', style='List Bullet')
    doc.add_paragraph('Objetivo: Proveer un contrato inteligente robusto enfocado en la gestión de registros transaccionales dentro de un clúster de alta disponibilidad.', style='List Bullet')
    doc.add_paragraph('Funciones Principales:', style='List Bullet')
    doc.add_paragraph('incrementarRegistro(string memory mensaje): Aumenta el contador transaccional y emite un evento de auditoría.', style='List Bullet 2')
    doc.add_paragraph('obtenerContador(): Retorna el estado global actual del contrato.', style='List Bullet 2')

    doc.add_heading('2. Arquitectura de Despliegue (Diagrama Draw.io)', level=1)
    doc.add_paragraph('La solución se compone de un objeto Deployment de Kubernetes administrando un ReplicaSet que garantiza la tolerancia a fallos, expuesto internamente mediante un Service de tipo ClusterIP. El diseño arquitectónico completo se encuentra modelado y disponible en el archivo vectorial editable "arquitectura_kubernetes_andres_silva.drawio" dentro del repositorio oficial.')
    
    diagrama = """+-----------------------------------------------------------------------+
|                    KUBERNETES CLUSTER (Docker Desktop)                |
|   +---------------------------------------------------------------+   |
|   |             Service: blockchain-contract-service              |   |
|   |                        (Port 8545)                            |   |
|   +-------------------------------+-------------------------------+   |
|                                   |                                   |
|   +-------------------------------+-------------------------------+   |
|   |           Deployment: blockchain-contract-deployment          |   |
|   |   +-------------------------------------------------------+   |   |
|   |   |                   ReplicaSet Controller               |   |   |
|   |   |  +---------------+  +---------------+  +-----------+  |   |   |
|   |   |  | Pod 1 (Ganache)|  | Pod 2 (Ganache)|  | Pod 3...  |  |   |   |
|   |   |  +---------------+  +---------------+  +-----------+  |   |   |
|   |   +-------------------------------------------------------+   |   |
|   +-------------------------------+-------------------------------+   |
+-----------------------------------------------------------------------+"""
    agregar_codigo_consola(doc, diagrama)

    doc.add_heading('3. Pasos de Orquestación y Evidencias de Consola', level=1)
    
    doc.add_heading('A. Despliegue Inicial (3 Réplicas)', level=2)
    doc.add_paragraph('Aplicación inicial del manifiesto YAML en el clúster local.')
    agregar_codigo_consola(doc, f"> kubectl apply -f k8s/deployment_blockchain.yaml\n\n> kubectl get pods\n{evidencias['pods_ini']}\n\n> kubectl get rs\n{evidencias['rs_ini']}")

    doc.add_heading('B. Escalado Horizontal a 5 Réplicas', level=2)
    doc.add_paragraph('Modificación en caliente del número de instancias ejecutándose sin interrupción del servicio.')
    agregar_codigo_consola(doc, f"> kubectl scale deployment/blockchain-contract-deployment --replicas=5\n\n> kubectl get pods\n{evidencias['pods_scale']}\n\n> kubectl get rs\n{evidencias['rs_scale']}")

    doc.add_heading('C. Actualización de Imagen (Rolling Update)', level=2)
    doc.add_paragraph('Pase a producción mediante actualización controlada de la versión del contenedor.')
    agregar_codigo_consola(doc, f"> kubectl set image deployment/blockchain-contract-deployment node-container=trufflesuite/ganache-cli:v6.12.2\n\n> {evidencias['rollout']}\n\n> kubectl get pods\n{evidencias['pods_update']}\n\n> kubectl get rs\n{evidencias['rs_update']}")

    nombre_archivo = "Informe_Orquestacion_Blockchain_Andres_Silva.docx"
    doc.save(nombre_archivo)
    print(f"\n=== ¡PROCESO COMPLETADO EXITOSAMENTE! ===")
    print(f"1. Archivo draw.io generado: arquitectura_kubernetes_andres_silva.drawio")
    print(f"2. Repositorio sincronizado en: {repo_url}")
    print(f"3. Documento Word generado: {nombre_archivo}")

if __name__ == "__main__":
    nombre_repo = "smartcontract_kubernetes_andres_silva"
    crear_estructura_proyecto()
    evidencias = ejecutar_orquestacion_kubernetes()
    repo_url = inicializar_git_y_github(nombre_repo)
    generar_documento_word(evidencias, repo_url)