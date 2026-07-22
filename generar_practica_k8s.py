import os
import subprocess
import time
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Error: Falta la librería python-docx. Ejecuta: python -m pip install python-docx")
    exit()

def run_cmd(cmd):
    print("\n[EJECUTANDO] " + cmd)
    try:
        result = subprocess.run(cmd, shell=True, text=True, capture_output=True, check=True)
        return result.stdout.strip()
    except subprocess.CalledProcessError as e:
        return "Error: " + e.stderr

def agregar_codigo_consola(doc, texto):
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(12)

def insertar_imagen(doc, ruta_imagen, titulo):
    if os.path.exists(ruta_imagen):
        p = doc.add_paragraph()
        run = p.add_run()
        run.add_picture(ruta_imagen, width=Inches(6.0))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Leyenda de la imagen
        p_cap = doc.add_paragraph(titulo)
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.runs[0].italic = True
    else:
        doc.add_paragraph(f"[Espacio reservado para {ruta_imagen} - Imagen no encontrada en la carpeta]")

def generar_documento():
    print("\n=== INICIANDO AUTOMATIZACIÓN COMPLETA ===")
    
    doc = Document()
    
    # --- Portada ---
    titulo = doc.add_heading('Informe de Orquestación de Smart Contract en Kubernetes', 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('Autor: Andrés Alejandro Silva Aguilar').bold = True
    doc.add_paragraph('Institución: Universidad Técnica Particular de Loja (UTPL)')
    doc.add_paragraph('Módulo: Especialización en Arquitectura Cloud y Blockchain')
    doc.add_paragraph('Fecha: Julio 2026')
    
    p_git = doc.add_paragraph()
    p_git.add_run('Repositorio GitHub: ').bold = True
    p_git.add_run('https://github.com/aalejos3626/smartcontract_docker_Andres_Alejandro_Silva_Aguilar')
    
    # Incrustar captura del repositorio
    insertar_imagen(doc, "image_1a3fab.png", "Figura 1: Evidencia del repositorio GitHub actualizado.")
    
    doc.add_page_break()

    # --- Sección 1 y 2 ---
    doc.add_heading('1. Arquitectura de Despliegue', level=1)
    doc.add_paragraph('El sistema se orquesta bajo un Deployment de Kubernetes, que mantiene un ReplicaSet para garantizar la disponibilidad de los Pods. El tráfico interno se enruta mediante un Service (ClusterIP).')
    
    # Incrustar captura de Nodos y Servicios
    insertar_imagen(doc, "image_1a38c1.png", "Figura 2: Panel de Docker Desktop mostrando el Control Plane y el Servicio expuesto en el puerto 8545.")

    # --- Sección 3 ---
    doc.add_heading('2. Pasos Seguidos y Evidencias de Consola', level=1)
    
    print("\n--- Ejecutando Comandos en el Clúster ---")
    
    doc.add_heading('A. Cambio en Caliente 1: Escalado a 5 Réplicas', level=2)
    doc.add_paragraph('Se ejecutó el comando de escalado horizontal para aumentar la capacidad a 5 réplicas sin tiempo de inactividad.')
    run_cmd("kubectl scale deployment/smart-contract-deployment --replicas=5")
    time.sleep(8)
    out_pods2 = run_cmd("kubectl get pods")
    agregar_codigo_consola(doc, f"> kubectl scale deployment/smart-contract-deployment --replicas=5\n\n> kubectl get pods\n{out_pods2}")
    
    # Incrustar captura gráfica de los 5 pods
    insertar_imagen(doc, "image_1a384a.png", "Figura 3: Interfaz gráfica comprobando las 5 réplicas del deployment en estado Running.")

    doc.add_heading('B. Cambio en Caliente 2: Rolling Update', level=2)
    doc.add_paragraph('Se actualizó la imagen del contenedor simulando un pase a producción con una nueva versión del contrato inteligente.')
    run_cmd("kubectl set image deployment/smart-contract-deployment contract-container=trufflesuite/ganache-cli:v6.12.2")
    out_rollout = run_cmd("kubectl rollout status deployment/smart-contract-deployment")
    time.sleep(8)
    out_rs3 = run_cmd("kubectl get rs")
    agregar_codigo_consola(doc, f"> kubectl set image deployment/smart-contract-deployment contract-container=trufflesuite/ganache-cli:v6.12.2\n\n> kubectl rollout status deployment/smart-contract-deployment\n{out_rollout}\n\n> kubectl get rs\n{out_rs3}")

    # Guardar
    print("\n--- Generando Informe Word ---")
    nombre_archivo = "Informe_Orquestacion_Final.docx"
    doc.save(nombre_archivo)
    print(f"\n=== ¡PROCESO COMPLETADO! ===")
    print(f"El archivo '{nombre_archivo}' incluye tus comandos reales y capturas gráficas.")

if __name__ == "__main__":
    generar_documento()